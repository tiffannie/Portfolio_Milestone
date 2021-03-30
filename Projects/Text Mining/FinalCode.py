#---------------------------------------------------------------------------------
#----------------------------Generic Functions ------------------------------------
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import nltk
from sklearn.feature_selection import chi2
from sklearn import feature_selection
from sklearn import feature_extraction

import docx

PLOT_INDIVIDUAL_MODEL_GRAPHS = False

mydoc = docx.Document()

content = mydoc.add_paragraph("Summary of this analysis")
content.add_run("\nWord Counts")



rem_more =['ability', 'claim', 'consider', 'contract', 'link', 'unable', 'register', 'yahoo', 'please', 'address', 'associate', 'york', 'know', 'live', 'coverage', 'i', ',', 'x', '"', 'play', 'button', 'videos', 'de', 'la', 'dijo', 'it', 'el', 'image']
nltkstopwords = nltk.corpus.stopwords.words('english')
morestopwords = ['video', 'videos', 'play', 'button', 'read', 'click', 'more','know', 'https', 'http', 'well', 'said', 'one', 'time', 'people', 'look', 'many', 'ago', 'even', 'much', 'didnt', 'see', 'weve', 'say', 'ive', 'got', 'come', 'like', 'thats', 'ever', 'theyre', 'going', 'dont', 'want', 'rrthe', '\r', 'shall', 'made', 'et.', 'al', 'could','would','might','must','need',
                 'rrrrrr','rr', 'h','b', 'sha','wo','y',"'s","'d","'ll","'t","'m","'re","'ve", 
                 "n't", 'readingrrcoronavirus', 'treatmentrbut', 'recoveryrbookr22.95rview', "image", "reuters", "caption", "breaking", "news", "via", "via image caption", "copy", "copyright", "getty", "nbc", "cnn", "images", "using", "show", "result", "large", "also", "iv", "one", "two", "new", "previously", "shown"]

stopwords = nltkstopwords + morestopwords

#Data pre-processing
# filter out any tokens not containing letters (e.g., numeric tokens, raw punctuation)
def text_clean(text):
	text = text.replace("\n"," ").replace("\r"," ").replace('&amp', "").replace('br', "").replace('.',"")
	punctuations = '!#$%^&*()~!@<>/;:{}[]"?'
	t = str.maketrans(dict.fromkeys(punctuations, " "))
	text = text.translate(t)
	t = str.maketrans(dict.fromkeys("'`", ""))
	text = text.translate(t)
	return text

#Remove URL
def rem_url(text):
    text = re.sub(r'http\S+', '', text)
    return text

#Frequency distribution

def new_tfidf_v(X, y, n, m, max_tkns):
    tfidf = TfidfVectorizer(use_idf = True, ngram_range = (n,m), max_features = max_tkns)
    tfidf_m = tfidf.fit_transform(X)
    joblib.dump(tfidf,  'tfidf_v.pkl')
    df = pd.DataFrame(tfidf_m.toarray(), columns = tfidf.get_feature_names())
    df['Label'] = y    
    return df, tfidf_m, tfidf


def tfidf_chi_v(X,y, tfidf, tfidf_m, p_value):
    dtf_features = pd.DataFrame()
    for cat in np.unique(y):
        chi2, p = feature_selection.chi2(tfidf_m, y==cat)
        dtf_features = dtf_features.append(pd.DataFrame({"feature":tfidf.get_feature_names(), "score":1-p, "y":cat}))
        dtf_features = dtf_features.sort_values(["y","score"], ascending=[True,False])
        dtf_features = dtf_features[dtf_features["score"]> float(p_value)]
        X_names = dtf_features["feature"].unique().tolist()
    
    for cat in np.unique(y):
       print("# {}:".format(cat))
       print("  . selected features:", len(dtf_features[dtf_features["y"]==cat]))
       print("  . top features:", ",".join(dtf_features[dtf_features["y"]==cat]["feature"].values[:10]))
       print(" ")
    vectorizer = feature_extraction.text.TfidfVectorizer(vocabulary = X_names)
    tf = vectorizer.fit_transform(X)
    df = pd.DataFrame(tf.toarray(), columns = vectorizer.get_feature_names())
    df['Label'] = y    
    return df



from nltk import FreqDist

def n_grams(textwords, n):
    textdist = FreqDist(textwords)
    textitems = textdist.most_common(n)
    for item in textitems:
        print (item[0], '\t', item[1])


    #Lets look at the bigrams
    textdist = FreqDist((nltk.bigrams(textwords)))
    textitems = textdist.most_common(n)
    for item in textitems:
        print (item[0], '\t', item[1])
    
    #Lets look at the trigrams
    textdist = FreqDist((nltk.trigrams(textwords)))
    textitems = textdist.most_common(n)
    for item in textitems:
        print (item[0], '\t', item[1])
    

#count the words, count the docs, pass the corpus, remove words appearing less than n times and of less than k length
def word_counts(corpus, n, k):        
    #find total word count
    wordcount = {}
    
    for i in corpus:
        i = i.lower()
        for w in nltk.word_tokenize(i):
            if w not in wordcount.keys():
                wordcount[w] = 1
            else:
                wordcount[w] +=1
                
    #find word count by document.
    doccount ={}
    for ww in wordcount.keys():
        doccount[ww] = 0
    for i in corpus:
        i = i.lower()
        f = nltk.word_tokenize(i)
        for wordss in wordcount.keys():
            if wordss in f:
                doccount[wordss] += 1
    
    iter1dict = {}
    for key, value in doccount.items():
        if value > n:
            iter1dict[key] = value
    
    #remove terms that are smaller than 3 characters
    iter2dict = {}
    for key, value in iter1dict.items():
        if len(key) >k:
            iter2dict[key] = value
    
    cleanheadline = ""
    coll = iter2dict.keys()
    f = []
    for a in corpus:
        rtoken = nltk.word_tokenize(a)
        for w in rtoken:
            if w in iter2dict.keys():
                cleanheadline = cleanheadline + " " + w
        f.append(cleanheadline)
        cleanheadline = ""
    
    return f


from wordcloud import WordCloud, STOPWORDS, ImageColorGenerator
    
def word_cloud(text, i, stop_words):
    wordcloud = WordCloud(
                          background_color='white',
                          stopwords=stop_words,
                          max_words=100,
                          max_font_size=50, 
                          random_state=42
                         ).generate(str(text))
    print(wordcloud)
    fig = plt.figure(1)
    plt.figure(1,figsize=(13, 13))
    plt.imshow(wordcloud)
    plt.axis('off')
    plt.title('Wordcloud of key')
    plt.savefig('WordCloud_'+str(i)+'.png')
    plt.close()
    plt.show()

from sklearn import metrics
import seaborn as sns
def fit_model(classifier, X_train, y_train, X_test, y_test, model_name):
    y_array = pd.get_dummies(y_test, drop_first=False).values
    classes = y_test.unique()
    model = classifier.fit(X_train, y_train)
    predicted = model.predict(X_test)
    predicted_prob = model.predict_proba(X_test)
    accuracy = metrics.accuracy_score(y_test, predicted)
    auc = metrics.roc_auc_score(y_array, predicted_prob, 
                                multi_class="ovr")    
    print('Model Summary: '+ model_name+'\n==========================')
    print("Accuracy:",  round(accuracy,2))
    print("Auc:", round(auc,2))
    print("Detail:")
    print(metrics.classification_report(y_test, predicted))

    if PLOT_INDIVIDUAL_MODEL_GRAPHS:
        ## Plot confusion matrix
        cm = metrics.confusion_matrix(y_test, predicted)
        fig, ax = plt.subplots()
        sns.heatmap(cm, annot=True, fmt='d', ax=ax, cmap=plt.cm.Blues, 
                    cbar=False)
        ax.set(xlabel="Pred", ylabel="True", xticklabels=classes, 
               yticklabels=classes, title="Confusion matrix")
        plt.yticks(rotation=0)
    
        fig, ax = plt.subplots(nrows=1, ncols=2)
        ## Plot roc
        for i in range(len(classes)):
            fpr, tpr, thresholds = metrics.roc_curve(y_array[:,i],  
                                   predicted_prob[:,i])
            ax[0].plot(fpr, tpr, lw=3, 
                      label='{0} (area={1:0.2f})'.format(classes[i], 
                                      metrics.auc(fpr, tpr))
                       )
        ax[0].plot([0,1], [0,1], color='navy', lw=3, linestyle='--')
        ax[0].set(xlim=[-0.05,1.0], ylim=[0.0,1.05], 
                  xlabel='False Positive Rate', 
                  ylabel="True Positive Rate (Recall)", 
                  title="Receiver operating characteristic")
        ax[0].legend(loc="lower right")
        ax[0].grid(True)
        
        ## Plot precision-recall curve
        for i in range(len(classes)):
            precision, recall, thresholds = metrics.precision_recall_curve(
                         y_array[:,i], predicted_prob[:,i])
            ax[1].plot(recall, precision, lw=3, 
                       label='{0} (area={1:0.2f})'.format(classes[i], 
                                          metrics.auc(recall, precision))
                      )
        ax[1].set(xlim=[0.0,1.05], ylim=[0.0,1.05], xlabel='Recall', 
                  ylabel="Precision", title="Precision-Recall curve")
        ax[1].legend(loc="best")
        ax[1].grid(True)

        plt.show()

    return accuracy



from sklearn.feature_extraction.text import TfidfVectorizer
import joblib


def tfidf_v(X, y, n,m, max_tkns):
    tfidf = TfidfVectorizer(use_idf = True, ngram_range = (n,m), max_features = max_tkns)
    tfidf_m = tfidf.fit(X)
    tfidf_matrix = tfidf_m.transform(X)
    joblib.dump(tfidf,  'tfidf_v.pkl')
    df = pd.DataFrame(tfidf_matrix.toarray(), columns = tfidf.get_feature_names())
    df['Label'] = y
    return df

def test_tfidf_v(X,y):
    tfidf_mm = joblib.load('tfidf_v.pkl')
    tfidf_matrix = tfidf_mm.transform(X)
    df = pd.DataFrame(tfidf_matrix.toarray(), columns = tfidf_mm.get_feature_names())
    df['Label'] = y
    return df

    

def binary_v(X, y, n,m, max_tkns):
    tfidf = TfidfVectorizer(binary = True, use_idf = False, ngram_range = (n,m), max_features = max_tkns)
    tfidf_m = tfidf.fit(X)
    tfidf_matrix = tfidf_m.transform(X)
    joblib.dump(tfidf,  'binary_v.pkl')
    df = pd.DataFrame(tfidf_matrix.toarray(), columns = tfidf.get_feature_names())
    df['Label'] = y
    return df


def test_binary_v(X,y):
    tfidf_mm = joblib.load('binary_v.pkl')
    tfidf_matrix = tfidf_mm.transform(X)
    df = pd.DataFrame(tfidf_matrix.toarray(), columns = tfidf_mm.get_feature_names())
    df['Label'] = y
    return df

import matplotlib.pyplot as plt
def plt_feature_vec(tfidf_matrix_df):    
    X = tfidf_matrix_df.values.T
    plt.scatter(X[:, 0], X[:, 1])
    plt.axis('equal')




def print_grams(k):
    corpus = []
    for i in range(0, len(k)):    
        text = k[i]
        ##Convert to list from string
        text = text.split()
        text = " ".join(text)
        corpus.append(text)
    
    textwords = []
    for i in corpus:
        textwords.extend([word.lower() for sent in nltk.sent_tokenize(i) for word in nltk.word_tokenize(sent)])
    tok = pd.DataFrame({'words': textwords})
    print('Post cleaning, this corpus contains\n')
    print (str(tok.shape[0]) + ' Items')
    print('N-Grams')
    n_grams(textwords, 10)


import nltk
import matplotlib.pyplot as plt
import re
from nltk.stem import PorterStemmer
import gensim
import gensim.utils
from gensim.utils import lemmatize

#Tokenize the text
#Regular Expression
def regexp_tkn(text):
    wrd = re.compile(r'\w+')
    words = wrd.findall(text)
    return words

#Use PorterStemmer    
def stem_Porter(tkns_re):
    ps_Stemmer = PorterStemmer()
    stemmed_corpus_re_p = []
    for i in range(0, len(tkns_re)):
        k = [ps_Stemmer.stem(w) for w in tkns_re[i]]
        stemmed_corpus_re_p.append(k)
    corpus_ps = []
    sep = ' '
    len2 = []
    len3 = []
    len4 =[]
    len5 = []
    
    for i in stemmed_corpus_re_p:
        s = [w for w in i if len(w)>2]
        for w in i:
            if len(w) == 2:
                len2.append(w)
            elif len(w) == 3:
                len3.append(w)
            elif len(w) == 4:
                len4.append(w)
            else:
                len5.append(w)
        corpus_ps.append(sep.join(s))
    return corpus_ps

#Use GENSIM Lemmatizer
def lemma_GENSIM(tkns_re):
    stemmed_corpus_re_ge = []
    print('Gensim Lemmatization in progress..............')
    for i in range(0, len(tkns_re)):
        k = [lemmatize(w) for w in tkns_re[i]]
        stemmed_corpus_re_ge.append(k)
    
    #Split the gensim lemma into tokens and pos tag
    docs_tkns = []
    docs_pos = []
    tokens = []
    pos_tag = []
    for i in range(0, len(stemmed_corpus_re_ge)):
        f = stemmed_corpus_re_ge[i]
        for j in f:
            if len(j) != 0:
                d = str(j).split('/')
                tokens.append(d[0].replace("[b'", ''))
                pos_tag.append(d[1].replace("']", ''))
        docs_tkns.append(tokens)
        docs_pos.append(pos_tag)
        tokens = []
        pos_tag = []        
    #Detokenize and count by token size
    #Lemmatized Text
    corpus_lem = []
    sep = ' '
    len2 = []
    len3 = []
    len4 =[]
    len5 = []
    
    for i in docs_tkns:
        s = [w for w in i if len(w)>2]
        for w in i:
            if len(w) == 2:
                len2.append(w)
            elif len(w) == 3:
                len3.append(w)
            elif len(w) == 4:
                len4.append(w)
            else:
                len5.append(w)
        corpus_lem.append(sep.join(s))
    return corpus_lem


import matplotlib as mpl
from sklearn.manifold import MDS
def plot_cosine(dist, topics):
    MDS()
    mds = MDS(n_components=2, dissimilarity="precomputed", random_state=1)
    pos = mds.fit_transform(dist)  # shape (n_components, n_samples)
    xs, ys = pos[:, 0], pos[:, 1]
    print()
    print()
    
    topic_colors = {0: '#1b9e77', 1: '#d95f02', 2: '#FF334C', 3: '#33FF38', 4: '#4D33FF'}
    #set up cluster names using a dict
    topic_names = {0: 'business', 
                   1: 'sports', 
                   2: 'entertainment', 
                   3: 'politics',
                   4: 'medical'}
    #create data frame that has the result of the MDS plus the cluster numbers and titles
    df = pd.DataFrame(dict(x=xs, y=ys, label=topics) )
    #group by cluster
    groups = df.groupby('label')
    # set up plot
    fig, ax = plt.subplots(figsize=(17, 9)) # set size
    ax.margins(0.05) # Optional, just adds 5% padding to the autoscaling
    ax.set_title('Topic Clusters', size = 20)
    
    #iterate through groups to layer the plot
    #note that I use the cluster_name and cluster_color dicts with the 'name' lookup to return the appropriate color/label
    for name, group in groups:
        ax.plot(group.x, group.y, marker='o', linestyle='', ms=12, 
                label=topic_names[name], color=topic_colors[name], 
                mec='none')
        ax.set_aspect('auto')
        ax.tick_params(\
            axis= 'x',          # changes apply to the x-axis
            which='both',      # both major and minor ticks are affected
            bottom='off',      # ticks along the bottom edge are off
            top='off',         # ticks along the top edge are off
            labelbottom='off')
        ax.tick_params(\
            axis= 'y',         # changes apply to the y-axis
            which='both',      # both major and minor ticks are affected
            left='off',      # ticks along the bottom edge are off
            top='off',         # ticks along the top edge are off
            labelleft='off')
    ax.legend(numpoints=1)  #show legend with only 1 point
    plt.show() #show the plot

#---------------------- Plot the results ----------------------------
def plot_results(tfidf_acc_df, title):
    fig, ax = plt.subplots(figsize=(15, 10)) # set size
    ax.margins(0.05) # Optional, just adds 5% padding to the autoscaling
    ax.set_title(title, size = 20)
    ax.plot(tfidf_acc_df.Model, tfidf_acc_df.Accuracy, marker='o', linestyle='', ms=12)
    ax.set_aspect('auto')
    ax.tick_params(\
        axis= 'x',          # changes apply to the x-axis
        which='both',      # both major and minor ticks are affected
        bottom='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelbottom='off')
    ax.tick_params(\
        axis= 'y',         # changes apply to the y-axis
        which='both',      # both major and minor ticks are affected
        left='off',      # ticks along the bottom edge are off
        top='off',         # ticks along the top edge are off
        labelleft='on')
    ax.legend(numpoints=1)  #show legend with only 1 point
    plt.show() #show the plot
    

#----------------- Topic Modelling --------------------------------------
import gensim
from gensim.utils import simple_preprocess
from gensim.parsing.preprocessing import STOPWORDS
from nltk.stem import WordNetLemmatizer, SnowballStemmer
from nltk.stem.porter import *
import numpy as np
np.random.seed(2018)
import nltk
nltk.download('wordnet')
from gensim import corpora, models

def preprocess(text):
    result = []
    for token in gensim.utils.simple_preprocess(text):
        result.append(token)
    return result


#----------------------------------------- Collect data -------------------------
import requests
from bs4 import BeautifulSoup
import csv

corpus_file_path = 'D:\\Text Mining\\Data Files\\Updated_736_Code\\Final_Corpus.csv'
with open(corpus_file_path) as corpus_file:
    corpus = [tuple(line) for line in csv.reader(corpus_file)]

titles = []
articles = []
topics = []

# Remove header row.
corpus = corpus[1:]

count = 1
unsupported_count = 0
for label, url in corpus:
    print("Processing " + str(count) + " / " + str(len(corpus)) + " - " + str(label) + " : " + str(url))
    count += 1

    article = requests.get(url)
    article_content = article.content
    bs_article = BeautifulSoup(article_content, 'html5lib')
    if "bmj.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'highwire-cite-title') # BMJ
    elif "medsci.org" in url:
        article_title = bs_article.find_all('h1', class_ = 'title') # Medsci   
    elif "jci.org" in url:
        article_title = bs_article.find_all('h1', class_ = 'article-title') # jci 
    elif "abcnews.go.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'Article__Headline__Title') # ABC
    elif "cnn.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'pg-headline') # CNN
    elif "foxnews.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'headline') # Fox
    elif "foxbusiness.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'headline') # Fox
    elif "nbcnews.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'article-hero__headline') # NBC
    elif "nbcsports.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'entry-title') # NBC Sports
    elif "si.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'm-detail-header--title') # SI   
    elif "eonline.com" in url:
        article_title = bs_article.find_all('h1', class_ = 'article-detail__title') # E  
    elif "yahoo.com" in url:
        article_title = bs_article.find_all('h1') # Yahoo
    elif "bleacherreport.com" in url:
        article_title = bs_article.find_all('h1') # BR
    else:
        print("Unsupported website! Can't pull article title from url " + url + "!")
        unsupported_count += 1
        continue
    title = article_title[0].get_text()

    if "cnn.com" in url:
        body = bs_article.find_all('div', class_ = 'zn-body__paragraph') # CNN
    elif "nbcnews.com" in url:
        body = bs_article.find_all('p', class_ = 'endmarkEnabled') # NBC
    else:
        body = bs_article.find_all('p') # BMJ, Medsci, jci, ABC, Foc, SI, E
    x = []
    final_paragraph = []
    for k in range(0,len(body)):
        x.append(body[k].get_text())
        final_paragraph = " ".join(x)

    topics.append(label)
    titles.append(title)
    articles.append(final_paragraph)

print("\nFailed to process " + str(unsupported_count) + " URLs from the provided CSV.\n")

ds = {'Title': titles, 'Body': articles, 'Topic': topics}
final_df = pd.DataFrame(ds)
final_df.reset_index(drop = False, inplace = True)
final_df.drop(columns = 'index', inplace = True)

print(final_df)

final_df.to_csv('D:\\Text Mining\\Data Files\\Updated_736_Code\\final_df.csv', index = False)

#-------------------------------------------Cleanse and Model-------------------
#Standarize the categories

fig = plt.figure()
ax = plt.subplot(111)
ax.plot(final_df['Topic'].value_counts(), label='Document Count')
plt.title('Document count by category')
ax.legend()
#plt.show()

fig.savefig('plot.png')

#mydoc.add_picture('plot.png', width=docx.shared.Inches(5), height=docx.shared.Inches(7))

#mydoc.save('D:/Cod.docx')

#pres = {'Topic': {'business':'Finance', 'sports':'Sports', 'health': 'Medical'}}
pres = {'Topic': {'health': 'medical'}}
final_df.replace(pres, inplace = True)

#Drop null rows
final_df['Body'].isna().sum()
final_df = final_df[final_df['Body'].isna() == False]
final_df.reset_index(drop = True, inplace = True)

#final_df.drop(columns = ['index'], inplace = True)
len_headlines = []
for a in final_df['Body']:
    len_headlines.append(len(a))
final_df['Length'] = len_headlines

#Review histogram to see the distribution of news documents
pos_train = final_df['Body']
plt.hist(final_df['Length'])
plt.show()
print('There are total ' + str(len(pos_train))+ ' documents in the corpus')

chk_v = pos_train
textwords = []
for i in chk_v:
    i = nltk.word_tokenize(str(i))
    textwords.extend([word.lower() for word in i])

print('There are total ' + str(len(textwords))+ ' tokens in the corpus')
print('Top N Grams in the raw dataset')
n_grams(textwords, 10)

print('Lets clean the dataset')
#Step 1: Remove URL
d = []
for u in pos_train:
    d.append(rem_url(str(u)))

#Step 2: Remove punctuations and special characters
s = []
for u in d:
    s.append(text_clean(u))

#Check after cleansing
chk_v = s
textwords = []
cos = []
td = []
for i in chk_v:
    i = nltk.word_tokenize(str(i))
    cos = [word.lower() for word in i if not word in stopwords]
    textwords.extend([word.lower() for word in i if not word in stopwords])
    td.append(' '.join(cos))

print('There are total ' + str(len(textwords))+ ' tokens in the corpus after cleaning')
print('Top N Grams after cleansing')
n_grams(textwords, 10)

#Drop all the numbers from the text and more tokens that do not add value
#Drop all the non-english words
nltk.download('words')
words = set(nltk.corpus.words.words())
cos2 = []
td2 = []
flist = []
for i in td:
    d = nltk.word_tokenize(i)
    cos2 = [word for word in d if (word in words) and (not word in rem_more) and not word.isnumeric()]
    td2.append(' '.join(cos2))
    flist.extend([word for word in d if (word in words) and (not word in rem_more) and not word.isnumeric()])

print('Top N grams after removing all non-english words and numbers')
n_grams(flist, 10)

textwords = flist
tkns_re = []
for u in td2:
    tkns_re.append(regexp_tkn(u))    
#Append clean text to the dataframe

#Stemming the corpus
corpus = stem_Porter(tkns_re)
#remove words that appear in less than n documents.
f = word_counts(corpus, 3, 3)
final_df['Clean Text'] = f
print_grams(final_df['Clean Text'])

#Lemmatizing the corpus
corpus = lemma_GENSIM(tkns_re)
f = word_counts(corpus, 3, 3)
final_df['Lemmatized Text'] = f
print_grams(final_df['Lemmatized Text'])


#------------ Let's make wordclouds to visualize the topics ----------------
import matplotlib.pyplot as plt
cont = final_df['Lemmatized Text']
clss = final_df['Topic']
j = clss.unique()
names = [] 
sds = []
leng = []
arra = []
ffs = []
ww=0
for i in j:
    arra = []
    names.append('Cluster: '+str(i))
    leng.append(len(final_df[final_df['Topic']==i]['Lemmatized Text']))
    dff = final_df[final_df['Topic']==i]['Lemmatized Text']
    for yy in final_df[final_df['Topic']==i]['Lemmatized Text']:
        arra.append(yy)
    ffs.append(arra)
    word_cloud(arra, ww, stopwords)
    ww=ww+1

#----------- Lets model the topics -----------------------------------------
j=[]
topicdict = []
topiclust = []
for i in np.sort(final_df['Topic'].unique()):
    processed_docs = final_df[final_df['Topic']==i]['Lemmatized Text'].map(preprocess)
    dictionary = gensim.corpora.Dictionary(processed_docs)    
    #dictionary.filter_extremes(no_below=1, no_above=0.9, keep_n=100000)    
    bow_corpus = [dictionary.doc2bow(doc) for doc in processed_docs]    
    tfidf = models.TfidfModel(bow_corpus)


    #TFIDF for corpus
    corpus_tfidf = tfidf[bow_corpus]

    temtp = []
#Using TFIDF Corpus
    lda_model_tfidf = gensim.models.LdaMulticore(corpus_tfidf, num_topics=5, id2word=dictionary, passes=2, workers=4)
    #lda_model_tfidf = gensim.models.ldamodel.LdaModel(corpus_tfidf, num_topics=5, id2word=dictionary, passes=2)
    for idx, topic in lda_model_tfidf.print_topics(-1):
        for r in lda_model_tfidf.show_topic(idx, topn=10):
            temtp.append(r[0])
    j.append(i)
    topicdict.append(temtp)
    topiclust.append('Topic of Cluster'+str(i))
    
topicmodel = {'Cluster':topiclust, 'Topic': topicdict}
topicdf = pd.DataFrame(topicmodel)
from sklearn.metrics.pairwise import cosine_similarity

#----------------------- Vectorize the data first -----------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff = tfidf_v(X, y, 1, 3, 10000)

#-----------Visualize the data --------------------------------

topic_names = {'Label': {'business'      : 0, 
                         'sports'        : 1, 
                         'entertainment' : 2, 
                         'politics'      : 3,
                         'medical'       : 4}}

dff.replace(topic_names, inplace = True)
tfidf_matrix = dff.drop(columns = ['Label'])
dist = 1 - cosine_similarity(tfidf_matrix)

topics = dff['Label']
plot_cosine(dist, topics)

#----------------------------- Unsupervised model -----------------------------
#---------------------------- Vectorize--------------------------------------

X = final_df['Lemmatized Text']
y = final_df['Topic']
dff = tfidf_v(X, y, 1, 3, 10000)

#-----------Visualize the data --------------------------------
topic_names = {'Label': {'business'      : 0, 
                         'sports'        : 1, 
                         'entertainment' : 2, 
                         'politics'      : 3,
                         'medical'       : 4}}

dff.replace(topic_names, inplace = True)
tfidf_matrix = dff.drop(columns = ['Label'])
dist = 1 - cosine_similarity(tfidf_matrix)
#-------------------------------- KMeans ----------------------------------
from sklearn.cluster import KMeans
num_clusters = 5
km = KMeans(n_clusters=num_clusters)
km.fit(tfidf_matrix)
clusters = km.labels_.tolist()
final_df['Predicted Clusters'] = clusters
plot_cosine(dist, clusters)

#--------------------------- Hierarchichal Clustering ----------------------
from scipy.cluster.hierarchy import ward, dendrogram
linkage_matrix = ward(dist)

fig, ax = plt.subplots(figsize=(15, 10)) # set size
ax = dendrogram(linkage_matrix, orientation="right");
plt.tick_params(\
    axis= 'x',          # changes apply to the x-axis
    which='both',      # both major and minor ticks are affected
    bottom='off',      # ticks along the bottom edge are off
    top='off',         # ticks along the top edge are off
    labelbottom='off')

plt.tight_layout() #show plot with tight layout

#-------------------------- Wordclouds of Clusters ------------------------------------------
cont = final_df['Lemmatized Text']
clss = final_df['Predicted Clusters']
j = clss.unique()
names = [] 
sds = []
leng = []
arra = []
ffs = []
ww=0
for i in j:
    arra = []
    names.append('Cluster: '+str(i))
    leng.append(len(final_df[final_df['Predicted Clusters']==i]['Lemmatized Text']))
    dff = final_df[final_df['Predicted Clusters']==i]['Lemmatized Text']
    for yy in final_df[final_df['Predicted Clusters']==i]['Lemmatized Text']:
        arra.append(yy)
    ffs.append(arra)
    word_cloud(arra, ww, stopwords)
    ww=ww+1

#------------------------ Predicted Topics -------------------------------
j=[]
topicdict = []
topiclust = []
for i in np.sort(final_df['Predicted Clusters'].unique()):
    processed_docs = final_df[final_df['Predicted Clusters']==i]['Lemmatized Text'].map(preprocess)
    dictionary = gensim.corpora.Dictionary(processed_docs)    
    #dictionary.filter_extremes(no_below=1, no_above=0.9, keep_n=100000)    
    bow_corpus = [dictionary.doc2bow(doc) for doc in processed_docs]    
    tfidf = models.TfidfModel(bow_corpus)


    #TFIDF for corpus
    corpus_tfidf = tfidf[bow_corpus]

    temtp = []
#Using TFIDF Corpus
    lda_model_tfidf = gensim.models.LdaMulticore(corpus_tfidf, num_topics=5, id2word=dictionary, passes=2, workers=4)
    #lda_model_tfidf = gensim.models.ldamodel.LdaModel(corpus_tfidf, num_topics=5, id2word=dictionary, passes=2)
    for idx, topic in lda_model_tfidf.print_topics(-1):
        for r in lda_model_tfidf.show_topic(idx, topn=10):
            temtp.append(r[0])
    j.append(i)
    topicdict.append(temtp)
    topiclust.append('Topic of Cluster'+str(i))
    
topicmodel = {'Cluster':topiclust, 'Topic': topicdict}
topicdf_predicted = pd.DataFrame(topicmodel)

#------------------------------- Lets compare Actual Topics vs Predicted topics ---------
print(topicdf_predicted['Topic'])
print(topicdf['Topic'])


#------------------------------------ Download test data set from different journals -
#Politics, Entertainment, Medical and Sports articles can be downloaded using this script
test_titles = []
test_articles = []

list_of_articles = ['https://news.yahoo.com/coronavirus-updates-us-surpasses-11m-090012284.html', 
                    'https://news.yahoo.com/pfizer-vs-moderna-what-we-know-and-what-we-dont-212333998.html',
                    'https://news.yahoo.com/doctors-calling-quits-under-stress-193342087.html',
                    'https://www.yahoo.com/news/live-results-2020-election-day-trump-biden-050117825.html',
                    'https://www.yahoo.com/news/dark-winter-biden-says-lack-213637189.html',
                    'https://sports.yahoo.com/its-hard-not-to-look-at-the-texans-jack-easterby-and-understand-why-the-nfl-keeps-expanding-the-rooney-rule-224257978.html',
                    'https://sports.yahoo.com/baker-mayfield-told-nick-chubb-222926640.html',
                    'https://www.yahoo.com/entertainment/jingle-jangle-a-christmas-journey-inclusivity-instant-holiday-classic-013027172.html',
                    'https://www.yahoo.com/entertainment/dog-the-bounty-hunter-daughter-cecily-plans-to-wear-mom-beth-chapman-wedding-dress-201410931.html',
                    'https://www.yahoo.com/entertainment/wolfgang-van-halen-premieres-touching-musical-tribute-to-father-eddie-i-love-and-miss-you-pop-192352238.html']

for i in list_of_articles:
    r1 = requests.get(i)
    coverpage = r1.content
    soup1 = BeautifulSoup(coverpage, 'html5lib')
    coverpage_news = soup1.find_all('h1')
    title = coverpage_news[0].get_text()
    test_titles.append(title)
    article = requests.get(i)
    article_content = article.content
    soup_article = BeautifulSoup(article_content, 'html5lib')
    body = soup_article.find_all('p')
    x = []
    final_paragraph = []
    for k in range(0,len(body)):
        x.append(body[k].get_text())
        final_paragraph = " ".join(x)
    test_articles.append(final_paragraph)


ds_test = {'Title': test_titles, 'Body': test_articles, 'Topic': 'medical'}
new_df_test = pd.DataFrame(ds_test)
new_df_test['Topic'].iloc[3] = 'politics'
new_df_test['Topic'].iloc[4] = 'politics'

new_df_test['Topic'].iloc[5] = 'sports'
new_df_test['Topic'].iloc[6] = 'sports'

new_df_test['Topic'].iloc[7] = 'business'
new_df_test['Topic'].iloc[8] = 'entertainment'
new_df_test['Topic'].iloc[9] = 'entertainment'

new_df_test.reset_index(drop = False, inplace = True)
new_df_test.drop(columns = 'index', inplace = True)
new_df_test['Topic'].value_counts()

#------------------------------ Cleanse and Vectorize the test corpus --------------
#--------------- Use the best models from the above analysis and verify if the topic of the test
#------------------can be predicted ------------------------------------------------
pos_train = new_df_test['Body']
print('There are total ' + str(len(pos_train))+ ' documents in the corpus')

chk_v = pos_train
textwords = []
for i in chk_v:
    i = nltk.word_tokenize(str(i))
    textwords.extend([word.lower() for word in i])
print('There are total ' + str(len(textwords))+ ' tokens in the corpus')
print('Top N Grams in the raw dataset')
n_grams(textwords, 10)

print('Lets clean the dataset')
#Step 1: Remove URL
d = []
for u in pos_train:
    d.append(rem_url(str(u)))

#Step 2: Remove punctuations and special characters
s = []
for u in d:
    s.append(text_clean(u))

#Check after cleansing
chk_v = s
textwords = []
cos = []
td = []
for i in chk_v:
    i = nltk.word_tokenize(str(i))
    cos = [word.lower() for word in i if not word in stopwords]
    textwords.extend([word.lower() for word in i if not word in stopwords])
    td.append(' '.join(cos))

print('There are total ' + str(len(textwords))+ ' tokens in the corpus after cleaning')
print('Top N Grams after cleansing')
n_grams(textwords, 10)

#Drop all the numbers from the text and more tokens that do not add value
#Drop all the non-english words
nltk.download('words')
words = set(nltk.corpus.words.words())
cos2 = []
td2 = []
flist = []
for i in td:
    d = nltk.word_tokenize(i)
    #cos2 = [word for word in d if (word in words) and (not word in rem_more) and not word.isnumeric()]
    cos2 = [word for word in d if not word in rem_more and not word.isnumeric()]
    td2.append(' '.join(cos2))
    #flist.extend([word for word in d if (word in words) and (not word in rem_more) and not word.isnumeric()])
    flist.extend([word for word in d if not word in rem_more and not word.isnumeric()])

print('Top N grams after removing all non-english words and numbers')
n_grams(flist, 10)

textwords = flist
tkns_re = []
for u in td2:
    tkns_re.append(regexp_tkn(u))    
#Append clean text to the dataframe

#Stemmer results
corpus = stem_Porter(tkns_re)
#remove words that appear in less than n documents.
f = word_counts(corpus, 3, 3)
new_df_test['Clean Text'] = f
print_grams(new_df_test['Clean Text'])

#Lemmatized the text
corpus = lemma_GENSIM(tkns_re)
f = word_counts(corpus, 3, 3)
new_df_test['Lemmatized Text'] = f
print_grams(new_df_test['Lemmatized Text'])


from sklearn.model_selection import train_test_split
#------------------------------------------------------------------------------
#------------------------------ Supervised Model -------------------------------
#------------------------------------------------------------------------------
#-----------Accuracy dictionary -------------
accu = []
modl = []
sett = []

#---------------------- Vectorize ----------------------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff = tfidf_v(X, y, 1, 3, 100000)
terms = dff.columns[: -1]

#dff.describe()
#----------------------------------MNB------------------------------------
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
X = dff.drop(columns = ['Label'])
y = dff['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)
a = fit_model(classifier, X_train, y_train, X_test, y_test, 'MNB')

accu.append(a)
modl.append('MNB')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')


#------------------ BNB -----------------------------------------------------
from sklearn.naive_bayes import BernoulliNB
clf = BernoulliNB()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'BNB')

accu.append(a)
modl.append('BNB')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')


#-----------------------------kNN ------------------------------------------
from sklearn.neighbors import NearestNeighbors    
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import GridSearchCV
knn = KNeighborsClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'kNN')

accu.append(a)
modl.append('KNN')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')


#------------------------- SVM ---------------------------------------------
from sklearn import svm
clf = svm.SVC(probability = True, kernel = 'poly', gamma = 'scale', coef0 = 0.01)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'SVM')


accu.append(a)
modl.append('SVM - POLY')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')


#--------Linear Kernel ---------
svc = svm.SVC(kernel ='linear', C = 1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')


accu.append(a)
modl.append('SVM - LINEAR')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')

#------------- RBF -----------------------
clf = svm.SVC(C=1, gamma=0.1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')


accu.append(a)
modl.append('SVM - RBF')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')

#------------------ Decision Tree -----------
from sklearn.tree import DecisionTreeClassifier
clf = DecisionTreeClassifier()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('DECISION TREE')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')

#------------------ Random Forest  -----------
from sklearn.ensemble import RandomForestClassifier
clf = RandomForestClassifier(max_depth=4, random_state=0)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('RANDOM FOREST')
sett.append('GENSIM LEMMATIZER + TFIDF-(1,3)-MAX FEATURES')



dicto = {'Model': modl, 'Configuration': sett, 'Accuracy': accu}

tfidf_acc_df = pd.DataFrame(dicto)

plot_results(tfidf_acc_df, np.unique(sett)[0])

#-----------Accuracy dictionary -------------
accu = []
modl = []
sett = []

#------------------------------ Binary vectorization -------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff_b = binary_v(X, y, 1, 3, 100000)

X = dff_b.drop(columns = ['Label'])
y = dff_b['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)
#----------------------------------MNB------------------------------------
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accu.append(a)
modl.append('MNB')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#------------------ BNB -----------------------------------------------------
from sklearn.naive_bayes import BernoulliNB
clf = BernoulliNB()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'BNB')
accu.append(a)
modl.append('BNB')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#-----------------------------kNN ------------------------------------------
from sklearn.neighbors import NearestNeighbors    
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import GridSearchCV
knn = KNeighborsClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'kNN')
accu.append(a)
modl.append('KNN')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#------------------------- SVM ---------------------------------------------
from sklearn import svm
clf = svm.SVC(probability = True, kernel = 'poly', gamma = 'scale', coef0 = 0.01)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'SVM')

accu.append(a)
modl.append('SVM - POLY')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#--------Linear Kernel ---------
svc = svm.SVC(kernel ='linear', C = 1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')

accu.append(a)
modl.append('SVM - LINEAR')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#------------- RBF -----------------------
clf = svm.SVC(C=1, gamma=0.1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')

accu.append(a)
modl.append('SVM - RBF')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#------------------ Decision Tree -----------
from sklearn.tree import DecisionTreeClassifier
clf = DecisionTreeClassifier()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('DECISION TREE')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')

#------------------ Random Forest  -----------
from sklearn.ensemble import RandomForestClassifier
clf = RandomForestClassifier(max_depth=4, random_state=0)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('RANDOM FOREST')
sett.append('GENSIM LEMMATIZER + BINARY-(1,3)-MAX FEATURES')


dicto = {'Model': modl, 'Configuration': sett, 'Accuracy': accu}

binary_acc_df = pd.DataFrame(dicto)

plot_results(binary_acc_df, np.unique(sett)[0])

#------------------------------------------- Porter Stemmer ------------------------------------
#-----------Accuracy dictionary -------------
accu = []
modl = []
sett = []

#---------------------- Vectorize ----------------------------------------
X = final_df['Clean Text']
y = final_df['Topic']
dff = tfidf_v(X, y, 1, 3, 100000)

#dff.describe()
#----------------------------------MNB------------------------------------
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
X = dff.drop(columns = ['Label'])
y = dff['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)
a = fit_model(classifier, X_train, y_train, X_test, y_test, 'MNB')

accu.append(a)
modl.append('MNB')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


#------------------ BNB -----------------------------------------------------
from sklearn.naive_bayes import BernoulliNB
clf = BernoulliNB()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'BNB')

accu.append(a)
modl.append('BNB')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


#-----------------------------kNN ------------------------------------------
from sklearn.neighbors import NearestNeighbors    
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import GridSearchCV
knn = KNeighborsClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'kNN')

accu.append(a)
modl.append('KNN')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


#------------------------- SVM ---------------------------------------------
from sklearn import svm
clf = svm.SVC(probability = True, kernel = 'poly', gamma = 'scale', coef0 = 0.01)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'SVM')


accu.append(a)
modl.append('SVM - POLY')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


#--------Linear Kernel ---------
svc = svm.SVC(kernel ='linear', C = 1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')
accu.append(a)
modl.append('SVM - LINEAR')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')

#------------- RBF -----------------------
clf = svm.SVC(C=1, gamma=0.1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')
accu.append(a)
modl.append('SVM - RBF')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


#------------------ Decision Tree -----------
from sklearn.tree import DecisionTreeClassifier
clf = DecisionTreeClassifier()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')
accu.append(a)
modl.append('DECISION TREE')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')

#------------------ Random Forest  -----------
from sklearn.ensemble import RandomForestClassifier
clf = RandomForestClassifier(max_depth=4, random_state=0)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('RANDOM FOREST')
sett.append('PORTER STEMMER + TFIDF-(1,3)-MAX FEATURES')


dicto = {'Model': modl, 'Configuration': sett, 'Accuracy': accu}

tfidf_acc_df = pd.DataFrame(dicto)

plot_results(tfidf_acc_df, np.unique(sett)[0])

#-----------Accuracy dictionary -------------
accu = []
modl = []
sett = []

#------------------------------ Binary vectorization -------------------------
X = final_df['Clean Text']
y = final_df['Topic']
dff_b = binary_v(X, y, 1, 3, 100000)

X = dff_b.drop(columns = ['Label'])
y = dff_b['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)
#----------------------------------MNB------------------------------------
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accu.append(a)
modl.append('MNB')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

#------------------ BNB -----------------------------------------------------
from sklearn.naive_bayes import BernoulliNB
clf = BernoulliNB()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'BNB')
accu.append(a)
modl.append('BNB')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

#-----------------------------kNN ------------------------------------------
from sklearn.neighbors import NearestNeighbors    
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import GridSearchCV
knn = KNeighborsClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'kNN')
accu.append(a)
modl.append('KNN')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

#------------------------- SVM ---------------------------------------------

from sklearn import svm
clf = svm.SVC(probability = True, kernel = 'poly', gamma = 'scale', coef0 = 0.01)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'SVM')

accu.append(a)
modl.append('SVM - POLY')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

#--------Linear Kernel ---------
svc = svm.SVC(kernel ='linear', C = 1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')


accu.append(a)
modl.append('SVM - LINEAR')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

#------------- RBF -----------------------
clf = svm.SVC(C=1, gamma=0.1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')

accu.append(a)
modl.append('SVM - RBF')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')


#------------------ Decision Tree -----------
from sklearn.tree import DecisionTreeClassifier
clf = DecisionTreeClassifier()
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')
accu.append(a)
modl.append('DECISION TREE')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')


#------------------ Random Forest  -----------
from sklearn.ensemble import RandomForestClassifier
clf = RandomForestClassifier(max_depth=4, random_state=0)
a = fit_model(clf, X_train, y_train,X_test, y_test, 'DT')

accu.append(a)
modl.append('RANDOM FOREST')
sett.append('PORTER STEMMER + BINARY-(1,3)-MAX FEATURES')

dicto = {'Model': modl, 'Configuration': sett, 'Accuracy': accu}
binary_acc_df = pd.DataFrame(dicto)

plot_results(binary_acc_df, np.unique(sett)[0])

#------------------------------------ Now, lets predict -----------------------------
# For selected models lets run the accuracy check using test dataset

#---------------------- Vectorize - TFIDF  ----------------------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff = tfidf_v(X, y, 1, 3, 100000)
terms = dff.columns[: -1]
X_train = dff.drop(columns = ['Label'])
y_train = dff['Label']


#--------------- Independent test data -----------------------------
X = new_df_test['Lemmatized Text']
y = new_df_test['Topic']
dftest = test_tfidf_v(X, y)
X_test = dftest.drop(columns = ['Label'])
y_test = dftest['Label']


#---------------------- Vectorize - Binary ----------------------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff = binary_v(X, y, 1, 3, 100000)
terms = dff.columns[: -1]
X_train = dff.drop(columns = ['Label'])
y_train = dff['Label']

#--------------- Independent test data -----------------------------
X = new_df_test['Lemmatized Text']
y = new_df_test['Topic']
dftest = test_binary_v(X, y)
X_test = dftest.drop(columns = ['Label'])
y_test = dftest['Label']

#--------Linear Kernel ---------
svc = svm.SVC(kernel ='linear', C = 1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')


#------------- RBF -----------------------
clf = svm.SVC(C=1, gamma=0.1, probability = True)
a = fit_model(svc, X_train, y_train,X_test, y_test, 'SVM')

from sklearn.neighbors import NearestNeighbors    
from sklearn.neighbors import KNeighborsClassifier
from sklearn.model_selection import GridSearchCV
knn = KNeighborsClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'kNN')

#------------------ Decision Tree -----------
from sklearn.tree import DecisionTreeClassifier
knn = KNeighborsClassifier()
clf = DecisionTreeClassifier()
a = fit_model(knn, X_train, y_train,X_test, y_test, 'DT')

#------------------- For best performing model, lets perform optimal cross validation & select salient features
X = final_df['Lemmatized Text']
y = final_df['Topic']

from sklearn.model_selection import cross_val_score
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
dff_b, tfidf_m, tfidf = new_tfidf_v(X, y, 1, 3, 100000)
X = dff_b.drop(columns = ['Label'])
y = dff_b['Label']
cross_val_score(classifier, X, y, cv=5)

#----------------------- Feature Importance ------------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff_b, tfidf_m, tfidf = new_tfidf_v(X, y, 1, 3, 100000)
X = dff_b.drop(columns = ['Label'])
y = dff_b['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.01)

classifier.fit(X_train, y_train)

print(classifier.classes_)

def get_salient_words(nb_clf, vect, class_ind):
    words = vect.get_feature_names()
    zipped = list(zip(words, nb_clf.feature_log_prob_[class_ind]))
    sorted_zip = sorted(zipped, key=lambda t: t[1], reverse=True)
    return sorted_zip

accrs = []
business_top_1000 = get_salient_words(classifier, tfidf, 0)[:1000]
entertainment_top_1000 = get_salient_words(classifier, tfidf, 1)[:1000]
medical_top_1000 = get_salient_words(classifier, tfidf, 2)[:1000]
politics_top_1000 = get_salient_words(classifier, tfidf, 3)[:1000]
sports_top_1000 = get_salient_words(classifier, tfidf, 4)[:1000]

business_top_2000 = get_salient_words(classifier, tfidf, 0)[:2000]
entertainment_top_2000 = get_salient_words(classifier, tfidf, 1)[:2000]
medical_top_2000 = get_salient_words(classifier, tfidf, 2)[:2000]
politics_top_2000 = get_salient_words(classifier, tfidf, 3)[:2000]
sports_top_2000 = get_salient_words(classifier, tfidf, 4)[:2000]



business_top_3000 = get_salient_words(classifier, tfidf, 0)[:3000]
entertainment_top_3000 = get_salient_words(classifier, tfidf, 1)[:3000]
medical_top_3000 = get_salient_words(classifier, tfidf, 2)[:3000]
politics_top_3000 = get_salient_words(classifier, tfidf, 3)[:3000]
sports_top_3000 = get_salient_words(classifier, tfidf, 4)[:3000]



business_top_4000 = get_salient_words(classifier, tfidf, 0)[:4000]
entertainment_top_4000 = get_salient_words(classifier, tfidf, 1)[:4000]
medical_top_4000 = get_salient_words(classifier, tfidf, 2)[:4000]
politics_top_4000 = get_salient_words(classifier, tfidf, 3)[:4000]
sports_top_4000 = get_salient_words(classifier, tfidf, 4)[:4000]


business_top_5000 = get_salient_words(classifier, tfidf, 0)[:5000]
entertainment_top_5000 = get_salient_words(classifier, tfidf, 1)[:5000]
medical_top_5000 = get_salient_words(classifier, tfidf, 2)[:5000]
politics_top_5000 = get_salient_words(classifier, tfidf, 3)[:5000]
sports_top_5000 = get_salient_words(classifier, tfidf, 4)[:5000]

#------ Lets re-run our models by selecting the salient feature -----

accrs = []
flist = []
for s in business_top_1000:
    flist.append(s[0])
for s in entertainment_top_1000:
    flist.append(s[0])
for s in medical_top_1000:
    flist.append(s[0])
for s in politics_top_1000:
    flist.append(s[0])
for s in sports_top_1000: 
    flist.append(s[0])
    
flist = np.unique(flist)
print('Total of ' + str(len(flist)) + ' features')
    

X = final_df['Lemmatized Text']
y = final_df['Topic']

vectorizer = TfidfVectorizer(vocabulary = flist)
X_m = vectorizer.fit_transform(X)
X_train, X_test, y_train, y_test = train_test_split(X_m, y, test_size = 0.2)
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accrs.append(a_full)



flist = []
for s in business_top_2000:
    flist.append(s[0])
for s in entertainment_top_2000:
    flist.append(s[0])
for s in medical_top_2000:
    flist.append(s[0])
for s in politics_top_2000:
    flist.append(s[0])
for s in sports_top_2000: 
    flist.append(s[0])
    
flist = np.unique(flist)
print('Total of ' + str(len(flist)) + ' features')
    

X = final_df['Lemmatized Text']
y = final_df['Topic']

vectorizer = TfidfVectorizer(vocabulary = flist)
X_m = vectorizer.fit_transform(X)
X_train, X_test, y_train, y_test = train_test_split(X_m, y, test_size = 0.2)
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accrs.append(a_full)


flist = []
for s in business_top_3000:
    flist.append(s[0])
for s in entertainment_top_3000:
    flist.append(s[0])
for s in medical_top_3000:
    flist.append(s[0])
for s in politics_top_3000:
    flist.append(s[0])
for s in sports_top_3000: 
    flist.append(s[0])
    
flist = np.unique(flist)
print('Total of ' + str(len(flist)) + ' features')
    

X = final_df['Lemmatized Text']
y = final_df['Topic']

vectorizer = TfidfVectorizer(vocabulary = flist)
X_m = vectorizer.fit_transform(X)
X_train, X_test, y_train, y_test = train_test_split(X_m, y, test_size = 0.2)
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accrs.append(a_full)


flist = []
for s in business_top_4000:
    flist.append(s[0])
for s in entertainment_top_4000:
    flist.append(s[0])
for s in medical_top_4000:
    flist.append(s[0])
for s in politics_top_4000:
    flist.append(s[0])
for s in sports_top_4000: 
    flist.append(s[0])
    
flist = np.unique(flist)
print('Total of ' + str(len(flist)) + ' features')
    

X = final_df['Lemmatized Text']
y = final_df['Topic']

vectorizer = TfidfVectorizer(vocabulary = flist)
X_m = vectorizer.fit_transform(X)
X_train, X_test, y_train, y_test = train_test_split(X_m, y, test_size = 0.2)
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accrs.append(a_full)


flist = []
for s in business_top_5000:
    flist.append(s[0])
for s in entertainment_top_5000:
    flist.append(s[0])
for s in medical_top_5000:
    flist.append(s[0])
for s in politics_top_5000:
    flist.append(s[0])
for s in sports_top_5000: 
    flist.append(s[0])
    
flist = np.unique(flist)
print('Total of ' + str(len(flist)) + ' features')
    

X = final_df['Lemmatized Text']
y = final_df['Topic']

vectorizer = TfidfVectorizer(vocabulary = flist)
X_m = vectorizer.fit_transform(X)
X_train, X_test, y_train, y_test = train_test_split(X_m, y, test_size = 0.2)
from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
accrs.append(a_full)

plt.plot(accrs, marker = "*")

#------------------------------------ Feature Selection ---------------------
#----------------------------------MNB------------------------------------
X = final_df['Lemmatized Text']
y = final_df['Topic']
dff_b, tfidf_m, tfidf = tfidf_v(X, y, 1, 3, 100000)
X = dff_b.drop(columns = ['Label'])
y = dff_b['Label']
X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)

from sklearn.naive_bayes import MultinomialNB
classifier = MultinomialNB()
a_full = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')

#------------------- Filter features --------------------------------------

p_value = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.70, 0.8, 0.9]
acs = []
for i in p_value:
    X = final_df['Lemmatized Text']
    y = final_df['Topic']
    dff_b = tfidf_chi_v(X,y, tfidf, tfidf_m, i)
    X = dff_b.drop(columns = ['Label'])
    y = dff_b['Label']
    X_train, X_test, y_train, y_test = train_test_split(X, y, test_size = 0.2)
    from sklearn.naive_bayes import MultinomialNB
    classifier = MultinomialNB()
    a_chi = fit_model(classifier, X_train, y_train,X_test, y_test, 'MNB')
    print('Actual Accuracy:'+ str(a_full)+ '- New Accuracy: '+str(a_chi)+ 'p-value of:'+str(i))
    acs.append(a_chi)
    if a_chi > a_full:
        print('Improvement after filtering')
    else:
        print('No improvement')


plt.plot(acs, marker = "*")